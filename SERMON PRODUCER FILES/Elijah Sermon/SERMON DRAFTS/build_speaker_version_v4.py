from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(r"C:\Users\onyei\Projects\chukwuma-theology")
SERMON_MD = BASE / "Elijah Sermon" / "SERMON DRAFTS" / "When_the_Tank_Is_Not_Low_Version_3.extracted.md"
OUT_DOCX = BASE / "Elijah Sermon" / "SERMON DRAFTS" / "When_the_Tank_Is_Not_Low_Version_4_Speaker_Cues.docx"


ACCENT = RGBColor(161, 100, 35)
DEEP = RGBColor(39, 47, 52)
MUTED = RGBColor(97, 105, 110)
SLIDE_FILL = "F4E8CF"
CUE_FILL = "EAF1EA"
IMPROVE_FILL = "EEF2F6"


SLIDES = [
    (1, "1 Kings 19:1-18 / When the Tank Is Not Low. It's Empty.", "Use as the congregation settles and as you introduce the text."),
    (2, "Central Proposition", "Advance after the Scripture reading and before the personal narrative becomes the doorway into the sermon."),
    (3, "Chapter 19 is not about fire. It is about bread.", "Advance when you make the turn from the childhood Mount Carmel memory to the burden of chapter 19."),
    (4, "One bad message", "Advance immediately after the chapter 18 background and before naming the collapse after victory."),
    (5, "Point One: God meets us at empty, not after empty.", "Advance as the first major movement begins."),
    (6, "What God Does Not Do / What God Does", "Advance before the contrast list. Let the visual carry the contrast while you slow down."),
    (7, "The journey is too great for thee.", "Advance before reading the KJV line; pause after the line."),
    (8, "Sabbath as creation architecture of rest", "Advance at the Sabbath distinctive. Treat this as the theological deepening of Point One."),
    (9, "Point Two: The God who speaks in the silence.", "Advance as the second movement begins."),
    (10, "Wind / Earthquake / Fire: Not there.", "Advance as you describe the three phenomena."),
    (11, "God made noise... then chose silence.", "Advance at the quote; let it land before explaining."),
    (12, "What noise is making it impossible to hear the thin silence?", "Advance as a diagnostic question to the room."),
    (13, "Point Three: God commissions broken people.", "Advance as the third movement begins."),
    (14, "Grace does not require full restoration before it trusts you again.", "Advance before the assignment paragraph; this is the pastoral release."),
    (15, "There are still seven thousand.", "Advance before v. 18 and remnant theology."),
    (16, "The bread is already baked.", "Advance for the closing appeal and keep it up through prayer."),
]


INSERTIONS_BEFORE = {
    "**Introduction · Personal Narrative · 4--6 Minutes**": [
        ("slide", 1),
        ("cue", "OPENING PACE: Warm, conversational, unhurried. Let the congregation meet the ten-year-old boy before you ask them to meet Elijah."),
        ("slide", 2),
        ("improve", "IMPROVEMENT: After the proposition, say one bridge sentence in your own voice: \"If you have ever been functional on the outside and finished on the inside, this story is for you.\" Then begin the Bronx/Liverpool story."),
    ],
    "Chapter 19 is not about fire. It is about bread. It is about sleep. It": [
        ("slide", 3),
        ("cue", "EMPHASIS: Slow down on the contrasts: fire / bread, chariot race / sleep, spectacle / silence. This is the sermon's first memorable turn."),
    ],
    "> ***The distance between your greatest victory and your deepest": [
        ("slide", 4),
        ("cue", "PAUSE: After \"one bad message,\" stop for two beats. Let people supply their own message, call, email, diagnosis, or conversation."),
    ],
    "**Point One · 10--12 Minutes**": [
        ("slide", 5),
        ("cue", "TRANSITION: Lower the volume slightly. This point should feel pastoral before it feels exegetical."),
    ],
    "## B. What God Does Not Do": [
        ("slide", 6),
        ("cue", "DELIVERY: Read the negative list cleanly, with a tiny pause after each verb. Do not rush to defend God. Let the absence of rebuke surprise them."),
    ],
    "*\"The journey is too great for thee.\" (v. 7, KJV)*": [
        ("slide", 7),
        ("cue", "EMPHASIS: Read the KJV line slower than the surrounding material. Let \"too great\" carry tenderness, not defeat."),
    ],
    "> **SDA Distinctive --- Sabbath as Created Architecture**": [
        ("slide", 8),
        ("improve", "IMPROVEMENT: Make Sabbath feel like good news, not a denominational talking point. Stress: Sabbath is not a reward for productivity; it is God's protest against human exhaustion."),
    ],
    "**Point Two · 10--12 Minutes**": [
        ("slide", 9),
        ("cue", "TRANSITION: Create contrast with Point One. God first restores the body; then He quiets the soul enough to hear."),
    ],
    "Then God does something remarkable. He sends three phenomena --- wind,": [
        ("slide", 10),
        ("cue", "PACING: Name wind, earthquake, and fire with space between them. The point is not that these are unimpressive; the point is that God refused to be reduced to them."),
    ],
    "> ***God made noise to show Elijah He could. Then He chose silence to": [
        ("slide", 11),
        ("cue", "PAUSE: After the quote, let the room stay quiet. Do not rescue the silence too quickly."),
    ],
    "> ***What noise in your life right now is making it impossible to hear": [
        ("slide", 12),
        ("improve", "IMPROVEMENT: Make this question concrete. Mention notifications, outrage, fear, overwork, religious performance, and the inner monologue that never stops."),
    ],
    "**Point Three · 8--10 Minutes**": [
        ("slide", 13),
        ("cue", "TRANSITION: Do not make this sound like God is sending Elijah back to grind harder. The order matters: care, silence, then commission."),
    ],
    "Grace does not require full restoration before it trusts you again. This": [
        ("slide", 14),
        ("cue", "EMPHASIS: This is a release line for wounded servants. Say it once, pause, then say the next sentence softly."),
    ],
    "*\"Yet I have reserved seven thousand in Israel, all the knees which": [
        ("slide", 15),
        ("improve", "IMPROVEMENT: Frame remnant theology humbly. Say clearly: the remnant is not a badge for superiority; it is evidence that God preserves witnesses by grace."),
    ],
    "**Conclusion and Closing Appeal · 5--7 Minutes**": [
        ("slide", 16),
        ("cue", "CLOSING PACE: Slow the entire conclusion by 15-20 percent. Let the repeated contrasts become prayerful rather than rhetorical."),
    ],
    "If you are empty today --- and you know who you are --- I want to invite": [
        ("cue", "APPEAL: Look up. This is the moment to speak directly to the room, not to the page. Keep your tone invitational, not pressured."),
    ],
    "*Read slowly. Let it settle. Close with prayer.*": [
        ("cue", "FINAL PAUSE: Read 1 Kings 19:7, close the manuscript, wait three beats, then pray. Do not add another mini-sermon after the text."),
    ],
    "## Key Hebrew Terms": [
        ("cue", "REFERENCE ONLY: This appendix is for preacher preparation. Do not read it unless you need a term during delivery."),
    ],
}

INSERTIONS_CONTAINS_BEFORE = {
    "The journey is too great for thee": [
        ("slide", 7),
        ("cue", "EMPHASIS: Read the KJV line slower than the surrounding material. Let \"too great\" carry tenderness, not defeat."),
    ],
    "Yet I have reserved seven thousand in Israel": [
        ("slide", 15),
        ("improve", "IMPROVEMENT: Frame remnant theology humbly. Say clearly: the remnant is not a badge for superiority; it is evidence that God preserves witnesses by grace."),
    ],
}


def clean_text(text: str) -> str:
    text = text.replace("\\'", "'").replace('\\"', '"')
    text = text.replace("---", "—")
    text = text.replace("--", "–")
    text = re.sub(r"\{\.mark\}", "", text)
    text = text.replace("\\.", ".")
    text = text.replace("\\", "")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def strip_md_inline(text: str) -> str:
    text = clean_text(text)
    text = re.sub(r"^\[|\]$", "", text)
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9C79F"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def add_runs_with_emphasis(paragraph, text: str):
    text = text.strip()
    if not text:
        return
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if not part:
            continue
        bold = False
        italic = False
        raw = part
        if raw.startswith("***") and raw.endswith("***"):
            bold = True
            italic = True
            raw = raw[3:-3]
        elif raw.startswith("**") and raw.endswith("**"):
            bold = True
            raw = raw[2:-2]
        elif raw.startswith("*") and raw.endswith("*"):
            italic = True
            raw = raw[1:-1]
        run = paragraph.add_run(clean_text(raw))
        run.bold = bold
        run.italic = italic


def add_callout(doc: Document, label: str, body: str, fill: str, label_color=ACCENT):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Inches(0.02)
    r = p.add_run(label + ": ")
    r.bold = True
    r.font.color.rgb = label_color
    r.font.size = Pt(9.5)
    body_run = p.add_run(body)
    body_run.font.size = Pt(9.5)
    body_run.font.color.rgb = DEEP
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_slide_callout(doc: Document, slide_no: int):
    title, note = next((title, note) for no, title, note in SLIDES if no == slide_no)
    add_callout(doc, f"ADVANCE SLIDE {slide_no}", f"{title} — {note}", SLIDE_FILL)


def add_note(doc: Document, note_type: str, text: str):
    if note_type == "cue":
        add_callout(doc, "SPEAKER CUE", text, CUE_FILL, RGBColor(59, 105, 72))
    else:
        add_callout(doc, "IMPROVEMENT", text, IMPROVE_FILL, RGBColor(63, 84, 115))


def configure_document(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.6)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    styles["Normal"].paragraph_format.line_spacing = 1.08

    for name, size, color in [
        ("Heading 1", 17, ACCENT),
        ("Heading 2", 13, DEEP),
        ("Heading 3", 11, MUTED),
    ]:
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_with_next = True


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("WHEN THE TANK IS NOT LOW.")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(24)
    run.font.color.rgb = DEEP
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("IT'S EMPTY.")
    r2.bold = True
    r2.font.name = "Aptos Display"
    r2.font.size = Pt(24)
    r2.font.color.rgb = ACCENT
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Version 4 Speaker Manuscript with Slide Prompts and Delivery Cues")
    sr.italic = True
    sr.font.size = Pt(12)
    sr.font.color.rgb = MUTED

    meta = [
        ("Text", "1 Kings 19:1-18"),
        ("Occasion", "Sabbath Morning Worship"),
        ("Congregation", "Braselton Seventh-day Adventist Church"),
        ("Speaker", "Dr. Chukwuma Onyeije"),
        ("Estimated Time", "38-45 minutes"),
        ("Tone Arc", "Warm personal -> exegetical -> pastoral -> direct appeal"),
    ]
    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate(meta):
        table.cell(i, 0).text = k
        table.cell(i, 1).text = v
        for cell in table.rows[i].cells:
            set_cell_border(cell, "E0E0E0")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(9.5)
        set_cell_shading(table.cell(i, 0), "F7F3EA")
        table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    add_callout(
        doc,
        "CENTRAL PROPOSITION",
        "When you arrive at empty, God's first response is not correction or commission; it is care. From that care, He restores and recommissions broken people.",
        "F7F3EA",
    )


def add_slide_map(doc: Document):
    doc.add_heading("Slide Cue Map", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ("Slide", "On-screen idea", "When to advance")
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = h
        set_cell_shading(cell, "27343A")
        set_cell_border(cell, "27343A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.8)
    for no, title, note in SLIDES:
        cells = table.add_row().cells
        cells[0].text = str(no)
        cells[1].text = title
        cells[2].text = note
        for cell in cells:
            set_cell_border(cell, "E1D8C8")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8.2)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_script(doc: Document):
    raw_lines = SERMON_MD.read_text(encoding="utf-8").splitlines()
    try:
        start = raw_lines.index("**Introduction · Personal Narrative · 4--6 Minutes**")
    except ValueError:
        start = 0
    lines = raw_lines[start:]

    doc.add_heading("Version 4 Speaker Script", level=1)

    pending_quote = []
    in_ascii_box = False
    ascii_box_lines = []
    inserted_contains = set()
    list_buffer = []

    def flush_quote():
        nonlocal pending_quote
        if pending_quote:
            text = " ".join(pending_quote)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.right_indent = Inches(0.18)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(7)
            r = p.add_run(strip_md_inline(text))
            r.italic = True
            r.font.color.rgb = RGBColor(75, 82, 86)
            pending_quote = []

    def flush_list():
        nonlocal list_buffer
        for item in list_buffer:
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_emphasis(p, item)
        list_buffer = []

    def flush_ascii_box():
        nonlocal ascii_box_lines
        cleaned = [strip_md_inline(x.strip("| ").strip()) for x in ascii_box_lines if x.strip("| ").strip()]
        cleaned = [x for x in cleaned if not re.match(r"^-+$", x)]
        if cleaned:
            add_callout(doc, "SCRIPTURE", " ".join(cleaned), "F7F3EA")
        ascii_box_lines = []

    for line in lines:
        if line in INSERTIONS_BEFORE:
            flush_quote()
            flush_list()
            for kind, payload in INSERTIONS_BEFORE[line]:
                if kind == "slide":
                    add_slide_callout(doc, payload)
                else:
                    add_note(doc, kind, payload)

        if any(line.startswith(marker) for marker in INSERTIONS_BEFORE if marker not in (line,)):
            for marker, items in INSERTIONS_BEFORE.items():
                if marker != line and line.startswith(marker):
                    flush_quote()
                    flush_list()
                    for kind, payload in items:
                        if kind == "slide":
                            add_slide_callout(doc, payload)
                        else:
                            add_note(doc, kind, payload)
                    break

        for marker, items in INSERTIONS_CONTAINS_BEFORE.items():
            if marker in line and marker not in inserted_contains:
                flush_quote()
                flush_list()
                for kind, payload in items:
                    if kind == "slide":
                        add_slide_callout(doc, payload)
                    else:
                        add_note(doc, kind, payload)
                inserted_contains.add(marker)

        stripped = line.strip()
        if not stripped:
            flush_quote()
            flush_list()
            continue
        if stripped.startswith("+:") or stripped.startswith("+---"):
            flush_quote()
            flush_list()
            if in_ascii_box:
                flush_ascii_box()
                in_ascii_box = False
            else:
                in_ascii_box = True
            continue
        if in_ascii_box:
            boxed = stripped.strip("| ").strip()
            if boxed:
                ascii_box_lines.append(boxed)
            continue
        if stripped == "## Key Hebrew Terms":
            flush_quote()
            flush_list()
            break
        if stripped.startswith("|") and stripped.endswith("|"):
            continue
        if re.match(r"^-{3,}$", stripped):
            continue
        if stripped.startswith(">"):
            pending_quote.append(stripped.lstrip("> ").strip())
            continue
        if stripped.startswith("- "):
            flush_quote()
            list_buffer.append(stripped[2:])
            continue

        flush_quote()
        flush_list()

        if stripped.startswith("# "):
            doc.add_heading(strip_md_inline(stripped[2:]), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(strip_md_inline(stripped[3:]), level=2)
        elif stripped.startswith("**Point") or stripped.startswith("**Conclusion") or stripped.startswith("**Introduction"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            r = p.add_run(strip_md_inline(stripped).upper())
            r.bold = True
            r.font.color.rgb = ACCENT
            r.font.size = Pt(10.5)
        else:
            p = doc.add_paragraph()
            add_runs_with_emphasis(p, stripped)

    flush_quote()
    flush_list()
    flush_ascii_box()
    add_preacher_reference(doc)


def add_preacher_reference(doc: Document):
    doc.add_page_break()
    doc.add_heading("Preacher Reference", level=1)
    add_note(doc, "cue", "REFERENCE ONLY: Keep this material nearby for confidence. In delivery, use only what serves the moment.")

    doc.add_heading("Key Hebrew Terms", level=2)
    terms = [
        ("rav", "1 Kings 19:4", "Enough; abundant, more than sufficient, past full measure. Elijah is not saying a little. He is saying this exceeds his outer limit."),
        ("qum / nus", "1 Kings 19:3", "Arose and ran; urgency, panic, flight. This is not a measured retreat but a break."),
        ("qol demamah daqah", "1 Kings 19:12", "The sound of thin silence. Almost the sound of no sound; a whisper at the edge of perception."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, h in enumerate(("Term", "Verse", "Speaker meaning")):
        cell = table.cell(0, idx)
        cell.text = h
        set_cell_shading(cell, "27343A")
        set_cell_border(cell, "27343A")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.8)
    for term, verse, meaning in terms:
        cells = table.add_row().cells
        cells[0].text = term
        cells[1].text = verse
        cells[2].text = meaning
        for cell in cells:
            set_cell_border(cell, "E1D8C8")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for para in cell.paragraphs:
                para.paragraph_format.space_after = Pt(0)
                for run in para.runs:
                    run.font.size = Pt(8.6)

    doc.add_heading("Pullquotes for Bulletin / Slides", level=2)
    quotes = [
        "The distance between your greatest victory and your deepest collapse can be one bad message.",
        "God's first response to your emptiness is not a new assignment. It is a meal.",
        "God made noise to show Elijah He could. Then He chose silence to show Elijah who He was.",
        "God commissions broken people. That has always been His method.",
        "The tank being empty is not the end of the story. It is where the story gets honest.",
    ]
    for quote in quotes:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(quote)


def add_footer(doc: Document):
    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = "When the Tank Is Not Low. It's Empty. | Version 4 Speaker Manuscript"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = MUTED


def main():
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_slide_map(doc)
    add_script(doc)
    add_footer(doc)
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
