from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

NAVY = RGBColor(0x1a, 0x2a, 0x4a)
GOLD = RGBColor(0xC8, 0x92, 0x2A)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
SLIDE_BG = RGBColor(0xE8, 0xEF, 0xF8)  # light blue-gray for slide content box

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def add_paragraph(doc, text='', size=11, bold=False, italic=False, color=None,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4,
                  left_indent=None):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_mixed_paragraph(doc, parts, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                        space_before=0, space_after=4, left_indent=None):
    """parts: list of (text, bold, italic, color)"""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    for text, bold, italic, color in parts:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_slide_header(doc, text):
    """Gold bold 16pt header with top border"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    # Add top border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '12')
    top.set(qn('w:space'), '4')
    top.set(qn('w:color'), 'C8922A')
    pBdr.append(top)
    pPr.append(pBdr)
    run = p.add_run(text)
    set_font(run, size=16, bold=True, color=GOLD)
    return p

def add_section_subheading(doc, text, space_before=6):
    p = add_paragraph(doc, text, size=13, bold=True, color=NAVY,
                      space_before=space_before, space_after=3)
    return p

def add_slide_box_start(doc, label="SLIDE CONTENT"):
    """Add a shaded box label for slide content"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    # Shade the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8EFF8')
    pPr.append(shd)
    run = p.add_run(f"[ {label} ]")
    set_font(run, size=9, bold=True, color=NAVY)
    return p

def add_shaded_paragraph(doc, text='', size=11, bold=False, italic=False, color=None,
                         left_indent=0.3, space_before=0, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.right_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'E8EFF8')
    pPr.append(shd)
    if text:
        run = p.add_run(text)
        set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_bullet(doc, text, level=0, size=11, italic=False, bold=False, shaded=False, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    if shaded:
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'E8EFF8')
        pPr.append(shd)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    return p

def add_numbered_item(doc, num, text, size=11, italic=False, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(f"{num}. {text}")
    set_font(run, size=size, italic=italic, color=DARK_GRAY)
    return p

def add_notes_box_header(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8E8')
    pPr.append(shd)
    run = p.add_run("TEACHER NOTES")
    set_font(run, size=10, bold=True, color=GOLD)
    return p

def add_page_break(doc):
    doc.add_page_break()

# ─── SLIDE 1: TITLE SLIDE ────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("SABBATH SCHOOL TEACHER'S LESSON")
set_font(run, size=12, bold=False, color=GOLD)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(10)
run = p.add_run("Repentance and Forgiveness:\nReturning to the Heart of God")
set_font(run, size=24, bold=True, color=NAVY)

# Memory text box
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'E8EFF8')
pPr.append(shd)
run = p.add_run('"If we confess our sins, He is faithful and just to forgive us our sins\nand to cleanse us from all unrighteousness."')
set_font(run, size=12, bold=False, italic=True, color=NAVY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(12)
pPr2 = p._p.get_or_add_pPr()
shd2 = OxmlElement('w:shd')
shd2.set(qn('w:val'), 'clear')
shd2.set(qn('w:color'), 'auto')
shd2.set(qn('w:fill'), 'E8EFF8')
pPr2.append(shd2)
run = p.add_run("— 1 John 1:9, NKJV")
set_font(run, size=11, bold=True, color=NAVY)

add_paragraph(doc, "Dr. Chukwuma Onyeije  |  Braselton Seventh-day Adventist Church",
              size=12, bold=True, color=NAVY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_paragraph(doc, "June 6, 2026  |  60 Minutes",
              size=11, color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_paragraph(doc, "Based on ALQ226 Lesson 10 by Dalibor Golubovic",
              size=11, italic=True, color=DARK_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

# Slide 1 label
add_paragraph(doc, "SLIDE 1 — TITLE SLIDE", size=9, bold=True, color=GOLD,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, space_before=8)

add_page_break(doc)

# ─── SLIDE 2: THE BIG IDEA ───────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 2 — THE BIG IDEA")

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc, "Three Core Theological Claims:", size=11, bold=True)
add_shaded_paragraph(doc, "1.  The journey from sin to restoration requires more than just feeling sorry; it demands genuine repentance.", size=11)
add_shaded_paragraph(doc, "2.  When we turn to God, we discover a Father who is eager to forgive.", size=11)
add_shaded_paragraph(doc, "3.  He offers us the priceless robe of Christ's righteousness to cover our inadequacies.", size=11, space_after=6)

add_notes_box_header(doc)
add_paragraph(doc, "Introduction — The Lesson Arc", size=12, bold=True, color=NAVY, space_before=4, space_after=2)
add_paragraph(doc,
    "This lesson traces a coherent theological journey: Conviction → Repentance → Forgiveness → Restoration. "
    "Five key passages guide us through each movement of this arc, showing how God's initiative meets human response "
    "at every stage.",
    size=11, color=DARK_GRAY, space_after=4)

add_paragraph(doc, "Five Key Passages:", size=11, bold=True, color=NAVY, space_before=2, space_after=2)
passages = [
    ("Isaiah 64:6", "— Contamination diagnosed"),
    ("Hosea 6:1–6", "— The healer identified"),
    ("Zechariah 3:1–5", "— The divine exchange performed"),
    ("Acts 3:18–19", "— The pathway walked daily"),
    ("Matthew 22:1–14", "— The warning heeded"),
]
for ref, desc in passages:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    run1 = p.add_run(ref)
    set_font(run1, size=11, bold=True, color=NAVY)
    run2 = p.add_run(desc)
    set_font(run2, size=11, color=DARK_GRAY)

add_paragraph(doc,
    "Orient class members to this arc at the outset so that each day's passage is understood as one movement "
    "in a larger symphony of grace.",
    size=11, italic=True, color=DARK_GRAY, space_before=4, space_after=4)

add_page_break(doc)

# ─── SLIDE 3: SUNDAY ─────────────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 3 — SUNDAY: THE RUSH OF LIFE")

add_paragraph(doc, "Key Text: Isaiah 64:6", size=12, bold=True, color=NAVY, space_before=2, space_after=4)

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc,
    "Busyness and the \"urgent\" distract us from the \"important\" — sitting at the feet of Jesus (Luke 10:41–42).",
    size=11)
add_shaded_paragraph(doc,
    "Recognizing our spiritual poverty allows us to exchange our \"filthy rags\" for Christ's pure robe "
    "of righteousness (Isa. 64:6).",
    size=11, space_after=6)

add_notes_box_header(doc)

add_paragraph(doc, "Opening Question", size=12, bold=True, color=NAVY, space_before=4, space_after=2)
add_paragraph(doc,
    '"Have you ever worked all day on something you were proud of, only to discover it wasn\'t what was needed?"',
    size=11, italic=True, color=DARK_GRAY, left_indent=0.3, space_after=6)

add_section_subheading(doc, "Exegetical Foundation — Hebrew Word Study: Beged Iddim")
items = [
    ("Literal meaning: ", "menstrual cloth, ceremonially unclean garments"),
    ("Cultural context: ", "Levitically disqualifying, untouchable under covenant law"),
    ("Theological significance: ", "even our best moral efforts are contaminated at the source"),
]
for label, detail in items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_font(r1, size=11, bold=True, italic=True, color=NAVY)
    r2 = p.add_run(detail)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "The Diagnostic Reality", space_before=6)
add_paragraph(doc,
    "Isaiah is describing the covenant community — people who knew God, had the temple, possessed the Law. "
    "Yet his diagnosis is devastating: not merely that their sins are unclean, but that their righteousnesses "
    "(plural) are contaminated. This is the starting point of authentic repentance: not cataloguing bad deeds, "
    "but recognizing that even good deeds, generated by unredeemed human will, are structurally flawed.",
    size=11, color=DARK_GRAY, space_after=4)

add_section_subheading(doc, "Connection to Sunday Theme")
add_paragraph(doc,
    "Like Martha distracted by \"urgent\" service while Mary chose the \"important\" (sitting at Jesus's feet), "
    "we can be so occupied with religious busyness that we never allow the Spirit to diagnose our actual condition.",
    size=11, color=DARK_GRAY, space_after=6)

add_paragraph(doc, "Discussion Questions  (5 min)", size=11, bold=True, color=NAVY, space_before=4, space_after=3)
dqs = [
    "Why does Isaiah describe even \"righteousnesses\" as contaminated? What does this teach about human nature?",
    "How does this verse prepare us for the need of divine intervention?",
    "In what areas might we be presenting \"contaminated\" righteousness to God?",
    "How does busyness prevent us from receiving the diagnosis that Isaiah 64:6 pronounces?",
]
for i, q in enumerate(dqs, 1):
    add_numbered_item(doc, i, q, italic=True)

add_page_break(doc)

# ─── SLIDE 4: MONDAY ─────────────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 4 — MONDAY: HOLY SPIRIT PROMPTINGS")

add_paragraph(doc, "Key Texts: Hosea 6:1–6 + John 16:8", size=12, bold=True, color=NAVY, space_before=2, space_after=4)

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc,
    "The Holy Spirit acts as a still, small voice, convicting us of sin (John 16:8).",
    size=11)
add_shaded_paragraph(doc,
    "We must resist the urge to justify our actions and let the Spirit lead us to the cross.",
    size=11, space_after=6)

add_notes_box_header(doc)

add_paragraph(doc, "Opening Bridge", size=12, bold=True, color=NAVY, space_before=4, space_after=2)
add_paragraph(doc,
    '"Isaiah shows us the problem — our contaminated condition. Now we must ask: what prompts us to address it? '
    'The answer is the wounded healer of Hosea 6."',
    size=11, italic=True, color=DARK_GRAY, left_indent=0.3, space_after=6)

add_section_subheading(doc, "Exegetical Foundation — Hosea 6:1–6")
add_paragraph(doc,
    "The Theological Shock: What's unexpected here is God's self-revelation: the same God who wounded is the God "
    "who heals. There's no appeal to a different, kinder deity. The God who disciplines is the God who restores.",
    size=11, color=DARK_GRAY, space_after=4)

add_section_subheading(doc, "Hebrew Precision", space_before=4)
hebrew_terms = [
    ('"He has torn us" (taraf)', ": divine discipline is real and necessary"),
    ('"He will heal us" (rapha)', ": the same hand that wounds provides healing"),
    ('"Like spring rains" (malkosh)', ": God's nature is reliable, seasonal, life-giving"),
]
for term, meaning in hebrew_terms:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(term)
    set_font(r1, size=11, bold=True, italic=True, color=NAVY)
    r2 = p.add_run(meaning)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "The Rain and Dew Contrast (v. 4)", space_before=6)
add_paragraph(doc,
    "Israel's love is like morning dew — evaporating under the first heat of testing. God's love is like spring "
    "rains — reliable, life-giving, seasonal but certain. The asymmetry is essential: our love is fragile and "
    "inconsistent; God's is steadfast and structural.",
    size=11, color=DARK_GRAY, space_after=4)

add_section_subheading(doc, "Verse 6 — The Heart Revelation: Hebrew Hesed", space_before=4)
add_paragraph(doc,
    "Steadfast love, covenant faithfulness, loyal love that keeps costly promises. This isn't sentimental kindness "
    "— it's the love that doesn't break when promises are expensive to keep. God desires hesed, not mere sacrifice "
    "— He wants covenant relationship, not religious transactions. The Holy Spirit's conviction is an expression "
    "of hesed: He wounds so we will return to the only one who can truly heal.",
    size=11, color=DARK_GRAY, space_after=6)

add_paragraph(doc, "Discussion Questions  (5 min)", size=11, bold=True, color=NAVY, space_before=4, space_after=3)
dqs4 = [
    "How do we reconcile God's love with His discipline? What does it mean that the same hand wounds and heals?",
    "Why does God compare Himself to rain and us to dew? What does this asymmetry teach about grace?",
    "What's the difference between offering God hesed (mercy) versus sacrifice?",
    "What does it look like to resist the Spirit's promptings vs. surrendering to them?",
]
for i, q in enumerate(dqs4, 1):
    add_numbered_item(doc, i, q, italic=True)

add_page_break(doc)

# ─── SLIDE 5: TUESDAY ────────────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 5 — TUESDAY: REAL REPENTANCE")

add_paragraph(doc, "Key Text: Acts 3:18–19", size=12, bold=True, color=NAVY, space_before=2, space_after=4)

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc,
    "True repentance involves sincere sorrow for sin and a deliberate choice to abandon it, leading to a "
    "transformed life.",
    size=11)
add_shaded_paragraph(doc,
    "It is the goodness and kindness of God that leads us to this genuine repentance (Rom. 2:4).",
    size=11, space_after=6)

add_notes_box_header(doc)

add_paragraph(doc, "Context Setting", size=12, bold=True, color=NAVY, space_before=4, space_after=2)
add_paragraph(doc,
    "Peter preaches after healing the lame man at Beautiful Gate. His audience had participated in crucifying "
    "the Messiah. His call isn't primarily about guilt management — it's about fundamental reorientation.",
    size=11, color=DARK_GRAY, space_after=6)

add_section_subheading(doc, "Greek Precision: Metanoeō")
greek_items = [
    ("Literal meaning: ", "change of mind, turning of the nous (seat of reason and intention)"),
    ("Critical distinction: ", "not remorse (looks backward) but repentance (faces forward)"),
    ("Growth principle: ", "you cannot grow in a direction you have not yet faced"),
]
for label, detail in greek_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_font(r1, size=11, bold=True, italic=True, color=NAVY)
    r2 = p.add_run(detail)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Kairoi Anapsyxeōs — \"Times of Refreshing\"", space_before=6)
add_paragraph(doc,
    "Imagery: cool breath after suffocating heat, soldier's rest after the battle line.",
    size=11, color=DARK_GRAY, space_after=2)
items5 = [
    "Personal: every genuine repentance produces soul refreshing",
    "Eschatological: connected to Christ's return and restoration of all things",
]
for item in items5:
    add_bullet(doc, item, size=11)

add_section_subheading(doc, "The Hinge Function", space_before=6)
add_paragraph(doc,
    "In salvation's economy, repentance is the hinge. Sins are \"blotted out\" — ledger wiped clean (sanctuary "
    "imagery). Growth cannot happen on a foundation that hasn't been cleared.",
    size=11, color=DARK_GRAY, space_after=4)
add_paragraph(doc,
    "The distinction between remorse and repentance is critical: Judas had remorse — it looked backward at the "
    "act and produced despair. Peter had repentance — it looked forward toward Christ and produced restoration. "
    "The direction of your sorrow determines its fruit.",
    size=11, color=DARK_GRAY, space_after=6)

add_paragraph(doc, "Discussion Questions  (5 min)", size=11, bold=True, color=NAVY, space_before=4, space_after=3)
dqs5 = [
    "What's the difference between remorse and repentance? How does this affect spiritual growth?",
    "How do individual and corporate \"times of refreshing\" connect?",
    "Why does Peter promise refreshing rather than just forgiveness?",
    "How does God's goodness (Rom. 2:4) function as the engine of genuine repentance?",
]
for i, q in enumerate(dqs5, 1):
    add_numbered_item(doc, i, q, italic=True)

add_page_break(doc)

# ─── SLIDE 6: WEDNESDAY ──────────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 6 — WEDNESDAY: SUFFICIENT GRACE")

add_paragraph(doc, "Key Texts: Zechariah 3:1–5 + Isaiah 61:10", size=12, bold=True, color=NAVY,
              space_before=2, space_after=4)

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc,
    "God's character is fundamentally \"merciful and gracious, slow to anger, and abounding in mercy\" (Exod. 34:6).",
    size=11)
add_shaded_paragraph(doc,
    "While the wages of sin is death, the free gift of God is eternal life through Jesus Christ, where grace "
    "abounds much more than sin (Rom. 5:20–21).",
    size=11, space_after=6)

add_notes_box_header(doc)

add_section_subheading(doc, "The Courtroom Scene — Zechariah 3:1–5", space_before=4)
court_items = [
    ("Setting: ", "heavenly sanctuary, divine courtroom"),
    ("Characters: ", "Joshua (high priest = Israel), Angel of the Lord (Christ), Satan (prosecutor)"),
    ("Drama: ", "the most visually compressed presentation of forensic justification in Scripture"),
]
for label, detail in court_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(detail)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Hebrew Word Study", space_before=6)
heb_items6 = [
    ("Tsow", " (filthy garments): excrement-soiled clothing, maximum contamination"),
    ("Machalatsowth", " (rich robes): festal garments, celebration clothing for holy entrance"),
]
for term, meaning in heb_items6:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(term)
    set_font(r1, size=11, bold=True, italic=True, color=NAVY)
    r2 = p.add_run(meaning)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "The Substitutionary Transaction", space_before=6)
add_paragraph(doc,
    "Critical observation: Satan's accusation is not answered by Joshua's defense. It's answered by what Christ "
    "does. The prosecutor is silenced not by the defendant's character but by the mediator's action. The command "
    "is immediate, decisive, unilateral: \"Take away the filthy garments.\" No extended argument. No period of "
    "probation. The contamination Isaiah described is removed by divine action, not human effort.",
    size=11, color=DARK_GRAY, space_after=4)

add_section_subheading(doc, "The Forensic Reality", space_before=4)
add_paragraph(doc,
    "Grace is not merely God tolerating us despite our sin. It is God actively clothing us in what we could "
    "never produce. The sufficient grace of Exodus 34:6 — God's self-disclosure as merciful, gracious, slow "
    "to anger — is what makes the Zechariah transaction possible.",
    size=11, color=DARK_GRAY, space_after=6)

add_section_subheading(doc, "From Contamination to Celebration — Isaiah 61:10")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.paragraph_format.space_after = Pt(4)
r1 = p.add_run("Hebrew Intensity: Sos Asis")
set_font(r1, size=11, bold=True, italic=True, color=NAVY)
r2 = p.add_run(" — doubled joy, emphatic, intensified form. Not calm resignation but exultant celebration. "
    "The response doesn't fit a clinical waiting room; it fits a wedding. The covenant broken by unfaithfulness "
    "(Isaiah 64) is restored in nuptial celebration language. The healing is not merely functional — it's relational.")
set_font(r2, size=11, color=DARK_GRAY)

add_paragraph(doc, "Discussion Questions  (5 min)", size=11, bold=True, color=NAVY, space_before=4, space_after=3)
dqs6 = [
    "Why was Joshua speechless before the exchange? What does this teach about self-defense before God?",
    "How does this scene help us understand what happens when we confess our sins?",
    "What role does Christ play that we cannot play for ourselves?",
    "How does the joy of Isaiah 61:10 flow from the grace displayed in Zechariah 3?",
]
for i, q in enumerate(dqs6, 1):
    add_numbered_item(doc, i, q, italic=True)

add_page_break(doc)

# ─── SLIDE 7: THURSDAY ───────────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 7 — THURSDAY: THE MOST EXPENSIVE ROBE")

add_paragraph(doc, "Key Text: Matthew 22:1–14", size=12, bold=True, color=NAVY, space_before=2, space_after=4)

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc,
    "The parable of the wedding banquet (Matt. 22:1–14) warns against relying on our own works.",
    size=11)
add_shaded_paragraph(doc,
    "We must actively choose to accept and wear Christ's perfect, unblemished character — the only acceptable "
    "covering before God.",
    size=11, space_after=6)

add_notes_box_header(doc)

add_section_subheading(doc, "The Parable Structure", space_before=4)
structure_items = [
    ("Movement 1 — Rejected invitation: ", "privilege ≠ acceptance"),
    ("Movement 2 — Open invitation: ", "\"both bad and good\" gathered"),
    ("Movement 3 — Wedding garment: ", "showing up ≠ acceptance"),
]
for label, detail in structure_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(detail)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Ancient Context", space_before=6)
add_paragraph(doc,
    "Wedding garments were provided by the host, not a test of wealth. The test was whether you received what "
    "the host provided versus insisting on your own clothing.",
    size=11, color=DARK_GRAY, space_after=4)

add_section_subheading(doc, "The Theological Warning", space_before=4)
add_paragraph(doc,
    "Most dangerous spiritual condition: accepting God's invitation while rejecting God's provision. "
    "Wanting kingdom benefits without submitting to the King's terms.",
    size=11, color=DARK_GRAY, space_after=6)

add_section_subheading(doc, "\"Many Called, Few Chosen\" (v. 14)")
called_items = [
    ("Klētoi", " (called): universal invitation"),
    ("Eklektoi", " (chosen): response-based selection, not arbitrary predestination"),
]
for term, meaning in called_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(term)
    set_font(r1, size=11, bold=True, italic=True, color=NAVY)
    r2 = p.add_run(meaning)
    set_font(r2, size=11, color=DARK_GRAY)
add_paragraph(doc,
    "The breadth of invitation doesn't eliminate the necessity of genuine response.",
    size=11, italic=True, color=DARK_GRAY, left_indent=0.4, space_before=2, space_after=4)

add_section_subheading(doc, "Eschatological Connection — Isaiah 61:10 + Revelation 19", space_before=4)
add_paragraph(doc,
    "The robe of righteousness Isaiah describes is both imputed covering and the accumulated, sanctified "
    "character of a life lived wearing Christ's righteousness. Revelation 19 returns to this imagery: the "
    "marriage supper of the Lamb, the bride arrayed in fine linen, \"which is the righteous acts of the "
    "saints.\" The garment that begins as imputed becomes expressed in a transformed life.",
    size=11, color=DARK_GRAY, space_after=6)

add_paragraph(doc, "Discussion Questions  (5 min)", size=11, bold=True, color=NAVY, space_before=4, space_after=3)
dqs7 = [
    "What's the difference between accepting God's invitation and wearing God's provision?",
    "How can someone be in the church but not properly clothed spiritually?",
    "What does it mean to receive Christ's righteousness daily?",
    "How does the \"most expensive robe\" become an act of grace rather than demand?",
]
for i, q in enumerate(dqs7, 1):
    add_numbered_item(doc, i, q, italic=True)

add_page_break(doc)

# ─── SLIDE 8: INSIDE STORY ───────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 8 — INSIDE STORY: VOICE SPEAKS AT A COOKING POT")

add_notes_box_header(doc)
add_paragraph(doc,
    "This is the ALQ Inside Story segment. Briefly narrate or summarize: Miriam, struggling with pain and "
    "sadness in Zambia, heard a divine voice prompting her to read the Bible. Her obedience led to immediate "
    "physical and spiritual healing. This story illustrates the transformative power of God's Word and His "
    "personal care.",
    size=11, color=DARK_GRAY, space_before=4, space_after=6)

add_paragraph(doc, "Application Bridge", size=12, bold=True, color=NAVY, space_before=2, space_after=3)
add_paragraph(doc,
    "Like Miriam, the Holy Spirit's promptings (Monday's lesson) come to us in unexpected places and moments. "
    "Obedience to that still, small voice is the first step of real repentance.",
    size=11, color=DARK_GRAY, space_after=4)

add_page_break(doc)

# ─── SLIDE 9: WEEKLY SYNTHESIS ───────────────────────────────────────────────
add_slide_header(doc, "SLIDE 9 — WEEKLY SYNTHESIS: KEY PRINCIPLES")

add_slide_box_start(doc, "SLIDE CONTENT")
add_shaded_paragraph(doc, "Key Principles:", size=11, bold=True)
add_shaded_paragraph(doc,
    "1.  Repentance is a deep, Spirit-led sorrow for sin, coupled with a turning away from it.",
    size=11)
add_shaded_paragraph(doc,
    "2.  Forgiveness is freely offered by a gracious God who desires our restoration.",
    size=11)
add_shaded_paragraph(doc,
    "3.  Our only hope for salvation is being completely clothed in the righteousness of Christ.",
    size=11, space_after=4)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(6)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'E8EFF8')
pPr.append(shd)
r1 = p.add_run("Memory Text: ")
set_font(r1, size=11, bold=True, color=NAVY)
r2 = p.add_run('"He restores my soul. He leads me in paths of righteousness for His name\'s sake." ')
set_font(r2, size=11, italic=True, color=DARK_GRAY)
r3 = p.add_run("— Psalm 23:3 (NKJV)")
set_font(r3, size=11, bold=True, color=NAVY)

add_notes_box_header(doc)

add_section_subheading(doc, "The Canonical Arc", space_before=4)
arc_items = [
    ("Isaiah 64:6", " (contamination diagnosed)"),
    ("Hosea 6:1–6", " (the healer identified)"),
    ("Zechariah 3:1–5", " (the exchange performed)"),
    ("Isaiah 61:10", " (the celebration entered)"),
    ("Acts 3:18–19", " (the pathway walked daily)"),
    ("Matthew 22:1–14", " (the warning heeded)"),
]
for ref, desc in arc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(ref)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(desc)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Memory Text Integration — 1 John 1:9", space_before=6)
add_paragraph(doc,
    "The God who is faithful and just is the same God who wounds and heals (Hosea 6), provides the divine "
    "exchange of garments (Zechariah 3), offers times of refreshing (Acts 3), and warns against self-provided "
    "covering (Matthew 22).",
    size=11, color=DARK_GRAY, space_after=6)

add_section_subheading(doc, "The Daily Clinical Protocol")
protocol_steps = [
    ("1. Diagnostic Honesty (Isaiah 64:6)", " — acknowledge contamination"),
    ("2. Approach the Wounded Healer (Hosea 6:1)", " — return to the One who disciplines in love"),
    ("3. Receive Divine Exchange (Zechariah 3:4)", " — accept Christ's righteousness"),
    ("4. Experience Wedding Joy (Isaiah 61:10)", " — celebrate relational restoration"),
    ("5. Enjoy Times of Refreshing (Acts 3:19)", " — rest in God's presence"),
    ("6. Wear God's Provision (Matthew 22:11)", " — daily put on Christ's righteousness"),
]
for label, detail in protocol_steps:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(detail)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Ellen White Illumination", space_before=6)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.4)
p.paragraph_format.space_after = Pt(4)
r1 = p.add_run('"Nothing is apparently more helpless, yet really more invincible, than the soul that feels '
    'its nothingness and relies wholly on the merits of the Savior."')
set_font(r1, size=11, italic=True, color=DARK_GRAY)
r2 = p.add_run(" — The Ministry of Healing, p. 182")
set_font(r2, size=11, bold=True, color=NAVY)

add_page_break(doc)

# ─── SLIDE 10: REFLECT, DISCUSS, APPLY ───────────────────────────────────────
add_slide_header(doc, "SLIDE 10 — REFLECT, DISCUSS, APPLY")

add_slide_box_start(doc, "SLIDE CONTENT — Discussion Questions")
dqs10_slide = [
    "How can we cultivate a deeper sensitivity to the Holy Spirit's promptings in our daily, busy lives?",
    "What is the difference between worldly sorrow over the consequences of sin and genuine, godly repentance?",
    "How does understanding Christ's robe of righteousness change the way we view our own spiritual efforts?",
]
for i, q in enumerate(dqs10_slide, 1):
    add_shaded_paragraph(doc, f"{i}.  {q}", size=11)

add_notes_box_header(doc)

add_section_subheading(doc, "Personal Reflection Questions", space_before=4)
reflection_qs = [
    "Where am I trying to present \"filthy rags\" as acceptable righteousness?",
    "How does understanding God as \"wounded healer\" change my approach to confession?",
    "What would daily \"putting on\" Christ's righteousness look like in my routine?",
    "How can I orient my desires toward God's truth rather than self-justification?",
]
for q in reflection_qs:
    add_bullet(doc, q, size=11)

add_section_subheading(doc, "Family / Community Application", space_before=6)
app_items = [
    ("Morning ritual: ", "begin each day with confession before performance"),
    ("Evening reflection: ", "review the day wearing Christ's righteousness versus self-effort"),
    ("Sabbath celebration: ", "gather in wedding joy rather than duty"),
    ("Community witness: ", "share the \"times of refreshing\" that follow genuine repentance"),
]
for label, detail in app_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(label)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(detail)
    set_font(r2, size=11, color=DARK_GRAY)

add_page_break(doc)

# ─── SLIDE 11: CLOSING ───────────────────────────────────────────────────────
add_slide_header(doc, "SLIDE 11 — CLOTHED IN GRACE (CLOSING)")

add_section_subheading(doc, "Memory Text Reprise — 1 John 1:9", space_before=4)
add_paragraph(doc, "The faithful and just God:", size=11, bold=True, color=NAVY, space_after=2)
closing_bullets = [
    "Tears and heals with the same hand",
    "Provides the garments He requires",
    "Offers refreshing, not just forgiveness",
    "Invites us to wedding celebration, not just relief",
]
for b in closing_bullets:
    add_bullet(doc, b, size=11)

add_section_subheading(doc, "Closing Prayer", space_before=8)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
p.paragraph_format.right_indent = Inches(0.3)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(10)
pPr = p._p.get_or_add_pPr()
shd = OxmlElement('w:shd')
shd.set(qn('w:val'), 'clear')
shd.set(qn('w:color'), 'auto')
shd.set(qn('w:fill'), 'FFF8E8')
pPr.append(shd)
run = p.add_run(
    '"Father, thank You that Your call to repentance is actually a call to refreshing. Help us daily wear '
    'what You provide rather than what we generate. May we approach You as the wounded healer who disciplines '
    'in love and restores in grace. Clothe us in Your righteousness so completely that when the enemy accuses, '
    'it is Your garments he sees. In the name of Jesus who stands as our mediator, Amen."')
set_font(run, size=11, italic=True, color=DARK_GRAY)

add_section_subheading(doc, "Benediction")
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.3)
p.paragraph_format.right_indent = Inches(0.3)
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(8)
run = p.add_run(
    '"May the Lord grant you a tender heart to hear His Spirit, the courage to turn from sin, and the profound '
    'peace of being fully covered by the righteousness of Jesus Christ. Amen."')
set_font(run, size=12, bold=True, italic=True, color=NAVY)

add_page_break(doc)

# ─── APPENDIX ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(8)
run = p.add_run("APPENDIX: TEACHER'S PREPARATION NOTES")
set_font(run, size=18, bold=True, color=NAVY)

# Hebrew / Greek Terms
add_section_subheading(doc, "Key Hebrew / Greek Terms for Reference")
terms = [
    ("Beged Iddim", " (Isa. 64:6): ", "filthy/contaminated garments — menstrual cloth"),
    ("Taraf", " (Hos. 6:1): ", "to tear/wound — divine discipline"),
    ("Rapha", " (Hos. 6:1): ", "to heal — divine restoration"),
    ("Malkosh", " (Hos. 6:3): ", "spring/latter rain — reliable, life-giving"),
    ("Hesed", " (Hos. 6:6): ", "steadfast love, covenant faithfulness"),
    ("Tsow", " (Zech. 3:3): ", "excrement-soiled — maximum contamination"),
    ("Machalatsowth", " (Zech. 3:4): ", "rich festal robes"),
    ("Sos Asis", " (Isa. 61:10): ", "doubled/intensified joy"),
    ("Metanoeō", " (Acts 3:19): ", "change of mind/direction — genuine repentance (not remorse)"),
    ("Kairoi Anapsyxeōs", " (Acts 3:19): ", "seasons/times of refreshing"),
    ("Klētoi", " (Matt. 22:14): ", "the called — universal invitation"),
    ("Eklektoi", " (Matt. 22:14): ", "the chosen — response-based"),
]
for term, ref, meaning in terms:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(term)
    set_font(r1, size=11, bold=True, italic=True, color=NAVY)
    r2 = p.add_run(ref)
    set_font(r2, size=11, bold=True, color=DARK_GRAY)
    r3 = p.add_run(meaning)
    set_font(r3, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Clinical Metaphors to Emphasize", space_before=8)
metaphors = [
    ("Contaminated self-treatment", " (Isaiah 64:6)"),
    ("Surgeon who wounds to heal", " (Hosea 6:1)"),
    ("Divine exchange in the treatment room", " (Zechariah 3:4)"),
    ("Surgical scrubs for a sterile environment", " (Matthew 22:11)"),
    ("Morning hospital preparation routine", " (1 John 1:9)"),
]
for m, ref in metaphors:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(m)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(ref)
    set_font(r2, size=11, italic=True, color=DARK_GRAY)

add_section_subheading(doc, "Cross-References for Extended Study", space_before=8)
cross_refs = [
    ("Revelation 19:8", " — Fine linen of the saints"),
    ("2 Corinthians 5:21", " — He became sin that we might become righteousness"),
    ("Galatians 3:27", " — Baptized into Christ, have put on Christ"),
    ("Romans 13:14", " — Put on the Lord Jesus Christ"),
    ("Psalm 51:6", " — You desire truth in the inner parts"),
]
for ref, desc in cross_refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(ref)
    set_font(r1, size=11, bold=True, color=NAVY)
    r2 = p.add_run(desc)
    set_font(r2, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Time Management Guide  (70 min total; trim discussions to fit 60 min)", space_before=8)
timing = [
    ("Slide 1 (Title/Opening):", "2 minutes"),
    ("Slide 2 (Big Idea):", "3 minutes"),
    ("Slide 3 (Sunday — Isaiah 64:6):", "10 minutes"),
    ("Slide 4 (Monday — Hosea 6:1–6):", "12 minutes"),
    ("Slide 5 (Tuesday — Acts 3:18–19):", "10 minutes"),
    ("Slide 6 (Wednesday — Zechariah 3 + Isaiah 61:10):", "10 minutes"),
    ("Slide 7 (Thursday — Matthew 22:1–14):", "10 minutes"),
    ("Slide 8 (Inside Story):", "3 minutes"),
    ("Slide 9 (Weekly Synthesis):", "5 minutes"),
    ("Slide 10 (Reflect, Discuss, Apply):", "3 minutes"),
    ("Slide 11 (Closing/Prayer):", "2 minutes"),
]
for label, time in timing:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.tab_stops.add_tab_stop(Inches(4.0))
    r1 = p.add_run(f"{label}\t{time}")
    set_font(r1, size=11, color=DARK_GRAY)

add_section_subheading(doc, "Interactive Elements", space_before=8)
interactive = [
    "Small group breakouts during discussion questions",
    "Personal reflection time during practical application",
    "Memory text recitation with Greek/Hebrew pronunciation",
    "Testimony sharing of \"times of refreshing\" experiences",
]
for item in interactive:
    add_bullet(doc, item, size=11)

# Save
output_path = "/home/user/chukwuma-theology/RepentanceAndForgivenessLesson_Revised.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
